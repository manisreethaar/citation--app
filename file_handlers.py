"""
file_handlers.py
-----------------
Read/write DOCX and PDF files, preserving as much formatting as possible.
"""

import os
import re
import sys
import tempfile
from pathlib import Path
from typing import Tuple


# ─── DOCX ────────────────────────────────────────────────────────────────────

def read_docx(path: str) -> Tuple[str, object]:
    """Return (full_text, Document_object)."""
    try:
        from docx import Document
    except ImportError:
        sys.exit("Missing dependency: pip install python-docx")

    doc = Document(path)
    paragraphs_text = []
    for para in doc.paragraphs:
        paragraphs_text.append(para.text)
    return '\n'.join(paragraphs_text), doc


def write_docx(original_doc, cited_body: str, cited_refs_text: str,
               refs_start_para_idx: int, output_path: str) -> None:
    """
    Write a new DOCX:
      - Replace paragraph text in body with cited version
      - Rebuild reference section with re-formatted bibliography
    """
    from docx import Document
    from docx.shared import Pt
    import copy

    # We do a paragraph-level replacement: split cited_body back into paragraphs
    cited_paras = cited_body.split('\n')

    doc = original_doc
    body_paras = doc.paragraphs[:refs_start_para_idx]

    for i, para in enumerate(body_paras):
        if i < len(cited_paras) and cited_paras[i] != para.text:
            # Clear runs and set new text, preserving first run's style
            if para.runs:
                for run in para.runs[1:]:
                    run.text = ''
                para.runs[0].text = cited_paras[i]
            else:
                para.text = cited_paras[i]

    # Replace everything from ref section onward with new bibliography
    # Remove old ref paragraphs
    for para in doc.paragraphs[refs_start_para_idx:]:
        p = para._element
        p.getparent().remove(p)

    # Append new bibliography
    for line in cited_refs_text.split('\n'):
        doc.add_paragraph(line)

    doc.save(output_path)


# ─── PDF ─────────────────────────────────────────────────────────────────────

def read_pdf(path: str) -> str:
    """Extract text from PDF. Returns full text string."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        sys.exit("Missing dependency: pip install PyMuPDF")

    doc = fitz.open(path)
    pages = []
    for page in doc:
        pages.append(page.get_text())
    return '\n'.join(pages)


def write_pdf_txt(cited_text: str, output_path: str) -> None:
    """
    Write cited text as a plain-text PDF using reportlab.
    Falls back to a .txt file if reportlab is unavailable.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm

        doc = SimpleDocTemplate(output_path, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        story = []
        for line in cited_text.split('\n'):
            if line.strip():
                story.append(Paragraph(line.replace('&', '&amp;'), styles['Normal']))
            else:
                story.append(Spacer(1, 0.3*cm))
        doc.build(story)
    except ImportError:
        # Fallback: save as txt
        txt_path = output_path.replace('.pdf', '_cited.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(cited_text)
        print(f"[Info] reportlab not installed — saved plain text to {txt_path}")


# ─── Generic text helpers ─────────────────────────────────────────────────────

def read_text(path: str) -> str:
    with open(path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def write_text(text: str, path: str) -> None:
    with open(path, 'w', encoding='utf-8') as f:
        f.write(text)


def find_refs_start_paragraph(doc, ref_section_re) -> int:
    """Return paragraph index where reference section starts."""
    import re
    for i, para in enumerate(doc.paragraphs):
        if ref_section_re.match(para.text):
            return i
    return len(doc.paragraphs)
