"""
file_io.py
===========
Read and write DOCX, PDF, and TXT files.
Clean abstraction — the pipeline only sees plain text.
"""

import os
import sys
from pathlib import Path


def read_file(path: str) -> str:
    """Read any supported file format and return plain text."""
    ext = Path(path).suffix.lower()
    if ext == '.docx':
        return _read_docx(path)
    elif ext == '.pdf':
        return _read_pdf(path)
    else:
        return _read_txt(path)


def write_file(text: str, output_path: str, original_path: str = None) -> None:
    """Write cited text to the output path."""
    ext = Path(output_path).suffix.lower()
    if ext == '.docx':
        _write_docx(text, output_path, original_path)
    elif ext == '.pdf':
        _write_pdf(text, output_path)
    else:
        _write_txt(text, output_path)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        sys.exit("Install python-docx:  pip install python-docx")

    doc = Document(path)
    return '\n'.join(p.text for p in doc.paragraphs)


def _write_docx(cited_text: str, output_path: str,
                original_path: str = None) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        sys.exit("Install python-docx:  pip install python-docx")

    # Start from original to preserve styles
    if original_path and Path(original_path).exists():
        doc = Document(original_path)
        orig_lines = [p.text for p in doc.paragraphs]
        new_lines = cited_text.split('\n')

        for i, para in enumerate(doc.paragraphs):
            if i < len(new_lines):
                new_text = new_lines[i]
                if new_text != para.text:
                    if para.runs:
                        para.runs[0].text = new_text
                        for run in para.runs[1:]:
                            run.text = ''
                    else:
                        para.text = new_text
    else:
        doc = Document()
        for line in cited_text.split('\n'):
            doc.add_paragraph(line)

    doc.save(output_path)


# ── PDF ───────────────────────────────────────────────────────────────────────

def _read_pdf(path: str) -> str:
    try:
        import fitz
    except ImportError:
        sys.exit("Install PyMuPDF:  pip install PyMuPDF")

    doc = fitz.open(path)
    return '\n'.join(page.get_text() for page in doc)


def _write_pdf(text: str, output_path: str) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm

        doc = SimpleDocTemplate(output_path, pagesize=A4,
                                leftMargin=2.5*cm, rightMargin=2.5*cm,
                                topMargin=2.5*cm, bottomMargin=2.5*cm)
        styles = getSampleStyleSheet()
        story = []
        for line in text.split('\n'):
            if line.strip():
                safe = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe, styles['Normal']))
            else:
                story.append(Spacer(1, 0.3*cm))
        doc.build(story)

    except ImportError:
        # Fallback: write as .txt alongside
        fallback = output_path.replace('.pdf', '_cited.txt')
        _write_txt(text, fallback)
        print(f"[Info] reportlab not installed — saved as {fallback}")


# ── TXT ───────────────────────────────────────────────────────────────────────

def _read_txt(path: str) -> str:
    with open(path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def _write_txt(text: str, path: str) -> None:
    with open(path, 'w', encoding='utf-8') as f:
        f.write(text)
